from __future__ import annotations

import datetime
import io

from docx import Document

from bahlily_export.docx_renderer import render_docx
from bahlily_export.models import ActionItem, ExportRequest, Quote


def _paragraph_texts(data: bytes) -> list[str]:
    doc = Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs]


def test_render_docx_includes_title_and_overview() -> None:
    req = ExportRequest(
        title="Team Standup",
        overview="Discussed sprint progress.",
        created_at=datetime.datetime(2026, 1, 1, tzinfo=datetime.UTC),
    )
    texts = _paragraph_texts(render_docx(req))
    assert "Team Standup" in texts
    assert "Discussed sprint progress." in texts


def test_render_docx_includes_generated_date() -> None:
    req = ExportRequest(
        title="T",
        overview="O",
        created_at=datetime.datetime(2026, 1, 1, tzinfo=datetime.UTC),
    )
    texts = _paragraph_texts(render_docx(req))
    assert any("2026-01-01" in t for t in texts)


def test_render_docx_key_points_as_bullets() -> None:
    req = ExportRequest(
        title="T",
        overview="O",
        key_points=["First point", "Second point"],
        created_at=datetime.datetime(2026, 1, 1, tzinfo=datetime.UTC),
    )
    doc = Document(io.BytesIO(render_docx(req)))
    bullet_paragraphs = [
        p for p in doc.paragraphs if p.style is not None and p.style.name == "List Bullet"
    ]
    bullet_texts = [p.text for p in bullet_paragraphs]
    assert "First point" in bullet_texts
    assert "Second point" in bullet_texts


def test_render_docx_action_item_with_owner_and_due_date() -> None:
    req = ExportRequest(
        title="T",
        overview="O",
        action_items=[ActionItem(description="Ship it", owner="Alex", due_date="2026-02-01")],
        created_at=datetime.datetime(2026, 1, 1, tzinfo=datetime.UTC),
    )
    texts = _paragraph_texts(render_docx(req))
    combined = " ".join(texts)
    assert "Ship it" in combined
    assert "Alex" in combined
    assert "2026-02-01" in combined


def test_render_docx_quote_attributed_to_speaker() -> None:
    req = ExportRequest(
        title="T",
        overview="O",
        quotes=[Quote(speaker="Alex", text="Let's ship it.", segment_id=3)],
        created_at=datetime.datetime(2026, 1, 1, tzinfo=datetime.UTC),
    )
    texts = _paragraph_texts(render_docx(req))
    combined = " ".join(texts)
    assert "Let's ship it." in combined
    assert "Alex" in combined
