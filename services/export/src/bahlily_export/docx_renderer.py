from __future__ import annotations

import io

from docx import Document

from bahlily_export.models import ExportRequest


def render_docx(req: ExportRequest) -> bytes:
    doc = Document()
    doc.add_heading(req.title, level=1)
    doc.add_paragraph(f"Generated on {req.created_at.strftime('%Y-%m-%d')}")

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(req.overview)

    if req.key_points:
        doc.add_heading("Key Points", level=2)
        for point in req.key_points:
            doc.add_paragraph(point, style="List Bullet")

    if req.action_items:
        doc.add_heading("Action Items", level=2)
        for item in req.action_items:
            text = item.description
            if item.owner:
                text += f" (Owner: {item.owner})"
            if item.due_date:
                text += f" (Due: {item.due_date})"
            doc.add_paragraph(text, style="List Bullet")

    if req.quotes:
        doc.add_heading("Quotes", level=2)
        for quote in req.quotes:
            speaker = quote.speaker or "Unknown"
            doc.add_paragraph(f'"{quote.text}" — {speaker}', style="Intense Quote")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
